import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import matplotlib.pyplot as plt
from matplotlib import dates as mpl_dates


def analysis_type_a(call_data_df, subscriber):
    df = call_data_df
    df['weigth_column'] = 1
    df['day_of_week'] = df['date_time'].dt.dayofweek
    df['hour_of_day'] = df['date_time'].dt.hour
    origin_type_count = df['type'].value_counts()  # origin types count (Series)
    original_types_list = origin_type_count.index.values.tolist()
    document = Document()
#---------------------------------------------------------------------------------------
#------------------------------------DOCUMENT SETTINGS----------------------------------
#---------------------------------------------------------------------------------------
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        section.bottim_margin = Cm(2)

    styles = document.styles
    # header style:
    main_header_style = styles.add_style('SBI_main_header', WD_STYLE_TYPE.PARAGRAPH)
    main_header_style.font.name = 'Times New Roman'
    main_header_style.font.size = Pt(14)
    main_header_style.font.bold = True
    main_header_style.paragraph_format.space_before = Pt(0)
    main_header_style.paragraph_format.space_after = Pt(0)
    main_header_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # body text style:
    body_text_style = styles.add_style('SBI_body_text', WD_STYLE_TYPE.PARAGRAPH)
    body_text_style.font.name = 'Times New Roman'
    body_text_style.font.size = Pt(14)
    body_text_style.paragraph_format.space_before = Pt(0)
    body_text_style.paragraph_format.space_after = Pt(0)
    body_text_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    body_text_style.paragraph_format.first_line_indent = Cm(1.25)
    body_text_style.paragraph_format.line_spacing = 1

    # red body text style:
    body_text_style_red = styles.add_style('SBI_body_text_red', WD_STYLE_TYPE.PARAGRAPH)
    body_text_style_red.font.name = 'Times New Roman'
    body_text_style_red.font.size = Pt(14)
    body_text_style_red.font.color.rgb = RGBColor(127, 12, 7)
    body_text_style_red.paragraph_format.space_before = Pt(0)
    body_text_style_red.paragraph_format.space_after = Pt(0)
    body_text_style_red.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    body_text_style_red.paragraph_format.first_line_indent = Cm(1.25)
    body_text_style_red.paragraph_format.line_spacing = 1

    # body text style:
    list_style = styles['List Bullet']
    list_style.font.name = 'Times New Roman'
    list_style.font.size = Pt(14)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(0)
    list_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    list_style.paragraph_format.first_line_indent = Cm(1.25)
    list_style.paragraph_format.line_spacing = 1

    # ---------------------------------------------------------------------------------------
    # ---------------------------------------HEADER MAIN-------------------------------------
    # ---------------------------------------------------------------------------------------
    # document title:
    first_header = document.add_paragraph()
    first_header.style = main_header_style
    first_header.add_run('А Н А Л І Т И Ч Н А   Д О В І Д К А').underline = True
    first_header_title1 = document.add_paragraph()
    first_header_title1.style = main_header_style
    first_header_title1.add_run("за результатами опрацювання технічних даних операторів рухомого ")
    first_header_title2 = document.add_paragraph()
    first_header_title2.style = main_header_style
    first_header_title2.add_run("(мобільного) зв'язку щодо абонента " + str(subscriber))
    space00 = document.add_paragraph()
    space00.style = body_text_style

    # ---------------------------------------------------------------------------------------
    # ---------------------------------------INTRO TEXT--------------------------------------
    # ---------------------------------------------------------------------------------------
    # get shape information:
    records_count = int(df.shape[0])  # num of all rows
    date_start_tab = df['date_time'].min()  # date of begin
    date_end_tab = df['date_time'].max()  # date of end
    plot_start = mpl_dates.date2num(date_start_tab)
    plot_end = mpl_dates.date2num(date_end_tab)
    tabs_day_lenth = (date_end_tab - date_start_tab).days + 1  # timedelta - investigation period

    # print intro text with shape info:
    intro_text = document.add_paragraph("Для складання аналітичної довідки надано відомості про з'єднання абонента ")
    intro_text.style = body_text_style
    intro_text.add_run(str(subscriber)).bold = True
    intro_text.add_run(' за період '+ str(tabs_day_lenth) + ' днів: '
                       + str(date_start_tab.strftime('%d.%m.%Y (з %H:%M год)'))
                       + ' - ' + str(date_end_tab.strftime('%d.%m.%Y (до %H:%M год)')) + '. ')
    intro_text.add_run('Всього записів: ' + str(records_count) + ', серед яких:')

    # ---------------------------------------------------------------------------------------
    # ---------------------------------------TYPES BULLET LIST-------------------------------
    # ---------------------------------------------------------------------------------------
    origin_type_list = list(df['type'].value_counts().index.values.tolist())  # list (like set) of virgin types

    # print types count:
    if 'вх' in origin_type_list:
         habit_list_types_1 = document.add_paragraph('записів про вхідні голосові дзвінки - '
                                                     + str(origin_type_count.at['вх']), style=list_style)
    if 'вих' in origin_type_list:
         habit_list_types_2 = document.add_paragraph('записів про вихідні голосові дзвінки - '
                                                     + str(origin_type_count.at['вих']), style=list_style)
    if 'вх СМС' in origin_type_list:
         habit_list_types_3 = document.add_paragraph('записів про вхідні СМС повідомлення - '
                                                     + str(origin_type_count.at['вх СМС']), style=list_style)
    if 'вих СМС' in origin_type_list:
         habit_list_types_4 = document.add_paragraph('записів про вихідні СМС повідомлення - '
                                                     + str(origin_type_count.at['вих СМС']), style=list_style)
    if 'переад' in origin_type_list:
         habit_list_types_5 = document.add_paragraph('переадресація (в т.ч. автовідповідач) - '
                                                     + str(origin_type_count.at['переад']), style=list_style)
    if 'internet' in origin_type_list:
         habit_list_types_6 = document.add_paragraph('записів про використання мережі інтернет - '
                                                     + str(origin_type_count.at['internet']), style=list_style)

    # ---------------------------------------------------------------------------------------
    # -----------------------------GETTING DATA FOR PLOTS------------------------------------
    #------------------------(detailed types and day duration)-------------------------------
    # ---------------------------------------------------------------------------------------
    # get info about non zero duration connections:
    non_null_dur_df = df.loc[df['dur'] > pd.Timedelta('0s')].copy()  # dataframe only non zero duration
    days_online = non_null_dur_df['date'].nunique()  # days with non zero duration
    days_intab = df['date'].nunique()  # days with all records


    # dict for plot labels:
    dict_types_desc = {'вх': 'вхідні дзвінки', 'вих': 'вихідні дзвінки', 'вх СМС': 'вхідні СМС',
                       'вих СМС': 'вихідні СМС', 'переад': 'переадресація',
                       'internet': 'використання мережі'}

    # create a coulumn with input connection ('0' second duration):
    def set_stat_type(row):
        if row['type'] == 'вх' and row['dur'] == pd.Timedelta('0s'):
            a = 'вхідні без відповіді'
            return a
        else:
            a = dict_types_desc.get(row['type'])  # change value to plot description
            return a

    df['types_stat'] = df.apply(lambda row: set_stat_type(row), axis=1)  # columns with plot description

    types_common_stat = df['types_stat'].value_counts()  # count types (in plot description)
    types_per_day = types_common_stat/tabs_day_lenth   # get types count in "per day" values with plot desc.

    total_dur_for_types = df.groupby(['type'])['dur'].sum()  # total dur of origin types
    dur_per_day_by_type = total_dur_for_types/days_online  # count dur per day (origin types)
    # set a string format (for a plot) origin types per day:
    dur_per_day_by_type = dur_per_day_by_type.apply(lambda x: f'{x.components.hours:02d}:{x.components.minutes:02d}:'
                                                              f'{x.components.seconds:02d} сек.')

    # get day dur for 'vh' 'vyh' types or fill "0":
    if 'вх' in dur_per_day_by_type:
        day_in_dur_str = dur_per_day_by_type.at['вх']
    else:
        day_in_dur_str = '00:00:00 cек.'

    if 'вих' in dur_per_day_by_type:
        day_out_dur_str = dur_per_day_by_type.at['вих']
    else:
        day_out_dur_str = '00:00:00 cек.'

    if 'вх' in total_dur_for_types:
        day_in_dur_val = total_dur_for_types.at['вх']
    else:
        day_in_dur_val = pd.Timedelta("0s")

    if 'вих' in total_dur_for_types:
        day_out_dur_val = total_dur_for_types.at['вих']
    else:
        day_out_dur_val = pd.Timedelta("0s")

    dur_day_val = (day_in_dur_val + day_out_dur_val)/days_online  # get duration of all voice connection per day
    all_time_voice_dur = day_in_dur_val + day_out_dur_val
    dur_day_str = f'{dur_day_val.components.hours:02d} год. {dur_day_val.components.minutes:02d} хв. ' \
                  f'{dur_day_val.components.seconds:02d} сек.'  # string of duration all voice connection per day

    # get list (of 2 types - 'in' and 'out' voice connection in 'int' format - for plot buiding):
    voice_types_int_list = [(day_out_dur_val/days_online).components.hours * 3600
                            + (day_out_dur_val/days_online).components.minutes * 60
                            + (day_out_dur_val/days_online).components.seconds,
                            (day_in_dur_val / days_online).components.hours * 3600
                            + (day_in_dur_val / days_online).components.minutes * 60
                            + (day_in_dur_val / days_online).components.seconds]

    # print info about active days:
    document.add_paragraph('', style=body_text_style)
    document.add_paragraph('Відомості щодо закономірностей у', style=main_header_style)
    document.add_paragraph('використанні абонентського номеру:', style=main_header_style)
    habit_day_active = document.add_paragraph('За період наданий для складання аналітичної довідки ('
                                              + str(tabs_day_lenth) + ' днів) абонент проявляв активність - '
                                              + str(days_online) + ' днів (мав розмови з іншими абонентами або '
                                                                   'користувався мережею інтернет). '
                                              'Використання абонентського номеру має наступні середні '
                                              'показники:', style=body_text_style)

    # ---------------------------------------------------------------------------------------
    # ---------------------------------------PLOTS MAKING------------------------------------
    # -------------------------(fig.1 - types per day & voice dur per day)-------------------
    # ---------------------------------------------------------------------------------------
    present_types = set(types_per_day.index.values.tolist())  # present type description
    if 'використання мережі' in present_types:
        types_per_day.drop(labels='використання мережі', inplace=True)  # drop internet statistic from h_bar

    #drawing types statistic plots:
    types_labels = types_per_day.index.values.tolist()
    plt.rcdefaults()
    plt.rcParams['font.size'] = 18.0
    fig1, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
    y_pos = np.arange(len(types_labels))
    ax1_prop = ax1.barh(types_labels, types_per_day, color='blue')
    ax1.set_yticks(y_pos)
    ax1.set(ylabel="ТИПИ З'ЄДНАНЬ")
    ax1.set(xlabel="КІЛЬКІСТЬ \n(в середньому на день)")
    ax1.invert_yaxis()
    #fig1.suptitle('В СЕРЕДНЬОМУ ЗА ДОБУ:', fontname="Times New Roman", fontsize=14, fontweight='bold')
    pie_type_list = ['Вихідні розмови:\n' + day_out_dur_str, 'Вхідні розмови:\n' + day_in_dur_str]
    ax2_prop = ax2.pie(voice_types_int_list, labels=pie_type_list, startangle=0, autopct='%0.1F%%', shadow=True,
                       wedgeprops={"edgecolor": "k", 'linewidth': 1}, labeldistance=1.2, colors=['cyan', 'magenta'],
                       explode=[0.1, 0])
    ax2.text(-1.25, -1.7, 'РОЗМОВ НА ДЕНЬ: ' + dur_day_str)
    plt.savefig("memfile.png", bbox_inches='tight')
    pic_01 = document.add_paragraph('', style=body_text_style)
    pic_1 = document.add_paragraph('', style=main_header_style)
    pic_1.add_run().add_picture('memfile.png', width=Cm(17))
    document.add_paragraph('', style=body_text_style)

    plt.close(fig1)
    # ---------------------------------------------------------------------------------------
    # -----------------------------TIMELINE PLOT - DURATION ---------------------------------
    #--------------------------(detailed duration sum dy dates)------------------------------
    # ---------------------------------------------------------------------------------------
    if dur_day_val > pd.Timedelta('0s'):
        df_voice_dur = df.loc[df['type'].isin(['вх', 'вих'])].copy()
        series_dur_voice = df_voice_dur.groupby(['date'])['dur'].sum()

        days_in_voice = int(series_dur_voice.size)
        values_voice_dur = series_dur_voice.astype('timedelta64[m]')
        values_voice_dur = values_voice_dur.tolist()
        date_voice_dur = series_dur_voice.index.values.tolist()
        for item in range(len(date_voice_dur)):
            date_voice_dur[item] = datetime.strptime(date_voice_dur[item], '%d.%m.%Y')
        date_voice_dur = mpl_dates.date2num(date_voice_dur)

        # get average total voice dur (for calc y lim plot):
        # 1.get total voice connection count:

        total_voice_dur = day_in_dur_val + day_out_dur_val
        total_voice_dur = total_voice_dur / np.timedelta64(1, 'm')
        total_voice_count = 0

        if 'вх' in origin_type_list:
            total_voice_count += int(origin_type_count.at['вх'])
        if 'вих' in origin_type_list:
            total_voice_count += int(origin_type_count.at['вих'])

        average_voice_dur = total_voice_dur/days_in_voice

        fig2, ax3 = plt.subplots()
        fig2.set_size_inches(18.5, 4, forward=True)
        ax3.xaxis.set_major_formatter(mpl_dates.DateFormatter("%d.%m.%Y"))
        plt.locator_params(axis='y', nbins=10)
        plt.locator_params(axis='x', nbins=25)
        ax3.set_xlim(plot_start, plot_end)
        ax3.set_ylim(0, average_voice_dur*3)
        plt.xticks(rotation=45)
        ax3.grid(True)
        ax3.set(xlabel="ДАТИ",
                ylabel="ТРИВАЛІСТЬ РОЗМОВ НА ДЕНЬ\n (хвилин)",
                title="ІНТЕНСИВНІСТЬ СПІЛКУВАННЯ У ПЕРІОДАХ ЧАСУ")
        ax3.bar(date_voice_dur, values_voice_dur, color='blue')
        plt.savefig("datafile.png", bbox_inches='tight')
        if all_time_voice_dur.components.days > 0:
            all_time_voice_dur_str = f'{all_time_voice_dur.components.days:02d} днів {all_time_voice_dur.components.hours:02d} ' \
                                     f'годин {all_time_voice_dur.components.minutes:02d} хвилин ' \
                                     f'{all_time_voice_dur.components.seconds:02d} секунд'
        else:
            all_time_voice_dur_str = f'{all_time_voice_dur.components.hours:02d} годин {all_time_voice_dur.components.minutes:02d} ' \
                                     f'хвилин {all_time_voice_dur.components.seconds:02d} секунд'
        pic_2 = document.add_paragraph('За час перевірки абонент провів розмов загалом на тривалість - '
                                       + all_time_voice_dur_str +
                                       "(сумарна тривалість вхідних та вихідних дзвінків). В розрізі періодів часу "
                                       "використання абонентського номеру, абонент спілкувався у періоди:",
                                       style=body_text_style)
        document.add_picture('datafile.png', width=Cm(17))
        document.add_paragraph('', style=body_text_style)
        plt.close(fig2)
    # ---------------------------------------------------------------------------------------
    # -----------------------------TIMELINE PLOT - INTERNET ---------------------------------
    #---------------------(detailed internet connection count dy dates)----------------------
    # ---------------------------------------------------------------------------------------
    if 'internet' in origin_type_list:

        df_internet_usage = df.loc[df['type'].isin(['internet'])].copy()
        series_internet_usage = df_internet_usage.groupby(['date'])['weigth_column'].sum()
        values_internet_usage = series_internet_usage.astype('int64')
        values_internet_usage = values_internet_usage.tolist()
        date_internet_usage = series_internet_usage.index.values.tolist()
        for item in range(len(date_internet_usage)):
            date_internet_usage[item] = datetime.strptime(date_internet_usage[item], '%d.%m.%Y')
        date_internet_usage = mpl_dates.date2num(date_internet_usage)

        fig3, ax4 = plt.subplots()
        fig3.set_size_inches(18.5, 4, forward=True)
        ax4.xaxis.set_major_formatter(mpl_dates.DateFormatter("%d.%m.%Y"))
        plt.locator_params(axis='y', nbins=10)
        plt.locator_params(axis='x', nbins=25)
        ax4.set_xlim(plot_start, plot_end)
        ax4.set_ylim(0, (origin_type_count.at['internet']/days_online)*3)
        plt.xticks(rotation=45)
        ax4.grid(True)
        ax4.set(xlabel="ДАТИ",
                ylabel="ВИКОРИСТАННЯ МЕРЕЖІ ІНТЕРНЕТ\n (кількість реєстрацій на день)",
                title="ІНТЕНСИВНІСТЬ ВИКОРИСТАННЯ МЕРЕЖІ У ПЕРІОДАХ ЧАСУ")
        ax4.bar(date_internet_usage, values_internet_usage, color='green')
        plt.savefig("datainternet.png", bbox_inches='tight')
        document.add_paragraph("Також абонент використовував мережу інтернет. Надається вибірка кількості "
                               "реєстрацій в мережі в розрізі періодів часу:", style=body_text_style)
        #document.add_paragraph('', style=body_text_style)
        document.add_picture('datainternet.png', width=Cm(17))
        document.add_paragraph('', style=body_text_style)
        plt.close(fig3)

    # ----------------------------------NULL PERIOD WARNING ---------------------------------
    date_points = pd.Series(non_null_dur_df['date_time'])
    date_points.sort_values()
    date_points = date_points.reset_index(drop=True)
    for i in range(date_points.size - 2):
        dif = date_points.at[i+1] - date_points.at[i]
        if dif > pd.Timedelta('280d'):
            start_null_period_str = date_points.at[i].strftime('%d.%m.%Y')
            end_null_period_str = date_points.at[i+1].strftime('%d.%m.%Y')
            null_period_lenth_str = str(dif.components.days)
            red_par = document.add_paragraph("Зважаючи на тривалу перерву у використанні абонентського номеру ("
                                   + start_null_period_str + " - " + end_null_period_str + ": всього "
                                   + null_period_lenth_str +" днів), наявна вірогідність перевипуску картки оператором "
                                                            "та придбання її новим власником.", style=body_text_style_red).bold = True
            break

    # ---------------------------------------------------------------------------------------
    # --------------------------BARS PLOT - HOUR OF A DAY USAGE -----------------------------
    #--------------------(detailed connection count by hour of a day)------------------------
    # ---------------------------------------------------------------------------------------

    if ("вх" in origin_type_list) or ("вих" in origin_type_list) or ("вих СМС" in origin_type_list):
        working_days = [0, 1, 2, 3, 4]
        rest_days = [5, 6]
        temp_hours_plot_types_list = ['вх', 'вих', 'вих СМС']

        # preparing a tab with WORKING days and active connections:
        df_voice_and_smsout_wd = df[df['type'].isin(temp_hours_plot_types_list) & df['day_of_week'].isin(working_days)].copy()
        df_voice_and_smsout_wd.drop(df_voice_and_smsout_wd[(df_voice_and_smsout_wd['dur'] == pd.Timedelta('0s'))
                                                           & (df_voice_and_smsout_wd['type'] == 'вх')].index,
                                    inplace=True)
        working_days_count = df_voice_and_smsout_wd['date'].nunique()

        # preparing a tab with REST days and active connections:
        df_voice_and_smsout_rd = df[df['type'].isin(temp_hours_plot_types_list) & df['day_of_week'].isin(rest_days)].copy()
        df_voice_and_smsout_rd.drop(df_voice_and_smsout_rd[(df_voice_and_smsout_rd['dur'] == pd.Timedelta('0s'))
                                                           & (df_voice_and_smsout_rd['type'] == 'вх')].index,
                                    inplace=True)

        rest_days_count = df_voice_and_smsout_rd['date'].nunique()


        #get Series of WORKING days stat by hour:
        con_by_hour_in_wd = df_voice_and_smsout_wd.groupby(['hour_of_day'])['weigth_column'].sum()
        con_by_hour_in_wd = con_by_hour_in_wd / working_days_count
        con_by_hour_in_wd.sort_index(inplace=True)
        con_by_hour_in_wd = con_by_hour_in_wd.reindex(range(24), fill_value=0)
        values_wd_by_hour = con_by_hour_in_wd.astype('float64')
        values_wd_by_hour = values_wd_by_hour.tolist()

        #get Series of REST days stat by hour:
        con_by_hour_in_rd = df_voice_and_smsout_rd.groupby(['hour_of_day'])['weigth_column'].sum()
        con_by_hour_in_rd = con_by_hour_in_rd / rest_days_count
        con_by_hour_in_rd.sort_index(inplace=True)
        con_by_hour_in_rd = con_by_hour_in_rd.reindex(range(24), fill_value=0)
        values_rd_by_hour = con_by_hour_in_rd.astype('float64')
        values_rd_by_hour = values_rd_by_hour.tolist()

        by_hour_x_ticks = ['00:00', '01:00', '02:00', '03:00', '04:00', '05:00', '06:00', '07:00', '08:00', '09:00',
                           '10:00', '11:00', '12:00', '13:00', '14:00', '15:00', '16:00', '17:00', '18:00', '19:00',
                           '20:00', '21:00', '22:00', '23:00']

        x = np.arange(len(by_hour_x_ticks))  # the label locations
        width = 0.3  # the width of the bars

        fig, ax = plt.subplots()
        rects1 = ax.bar(x - width / 2, values_wd_by_hour, width, label='У робочий день')
        rects2 = ax.bar(x + width / 2, values_rd_by_hour, width, label='У вихідний день')

        # Add some text for labels, title and custom x-axis tick labels, etc.
        ax.set_ylabel("З'ЄДНАНЬ НА ГОДИНУ")
        ax.set_title("СЕРЕДНЯ КІЛЬКІСТЬ З'ЄДНАНЬ У РОБОЧИЙ/ВИХІДНИЙ ДЕНЬ ПОГОДИННО\n ")
        ax.set_xticks(x)
        ax.set_xticklabels(by_hour_x_ticks)
        ax.legend()
        ax.grid(True)

        def autolabel(rects):
            """Attach a text label above each bar in *rects*, displaying its height."""
            for rect in rects:
                height = rect.get_height()
                ax.annotate('{:.1f}'.format(height),
                            xy=(rect.get_x() + rect.get_width() / 2, height),
                            xytext=(0, 3),  # 3 points vertical offset
                            textcoords="offset points",
                            ha='center', va='bottom', rotation='vertical')

        autolabel(rects1)
        autolabel(rects2)
        plt.xticks(rotation=45)
        fig.set_size_inches(15, 6, forward=True)
        fig.tight_layout()
        plt.savefig("by_hour.png", bbox_inches='tight')
        pic_3 = document.add_paragraph("Встановлені періоди доби в які абонент проявляв активність у спілкуванні. У "
                                       "вибірку вкюлчені записи про вихідні СМС повідомлення, вихідні дзвінки та "
                                       "вхідні дзвінки на які було надано відповідь. На графіку надається середня "
                                       "кількість записів про з'єднання для одного робочого дня та одного вихідного "
                                       "дня: ", style=body_text_style)
        document.add_paragraph('', style=body_text_style)
        document.add_picture('by_hour.png', width=Cm(17))
        document.add_paragraph('', style=body_text_style)
        plt.close(fig)

        #________________________________________NIGHT CALLS_-------------_______________________
        df_voice_and_smsout_night = df[df['hour_of_day'].isin([0, 1, 2, 3, 4, 5]) & df['type'].isin(['вх', 'вих', 'вих СМС'])].copy()
        if int(df_voice_and_smsout_night.shape[0]) != 0:
            night_abons_series = df_voice_and_smsout_night['sim_b'].value_counts()
            night_abons_series.sort_values(ascending=False, inplace=True)
            night_abons_list = night_abons_series.index.values.tolist()
            night_abons_values = night_abons_series.values.tolist()
            night_abons_lenth = len(night_abons_list)
            pic_3_5 = document.add_paragraph("Перевіркою вхідних, вихідних з'єднань, а також вихідних СМС повідомлень, "
                                             "що були здійснені в нічний час (у період з 00.00 год. до 06.00 год.), "
                                             "встановлено, що абонент мав нічні контакти з абонентами (вибірка за весь "
                                             "період перевірки): ", style=body_text_style)
            for i in range(night_abons_lenth):
                e = str(night_abons_values[i])
                if e.endswith('11') or  e.endswith('12') or e.endswith('13') or e.endswith('14'):
                    string = " з'єднань"
                elif e.endswith('1') or e.endswith('2') or e.endswith('3') or e.endswith('4'):
                    string = " з'єднання"
                else:
                    string = " з'єднань"
                document.add_paragraph(str(night_abons_values[i]) + string + " - абонент "
                                       + str(night_abons_list[i]), style=list_style)
        # ---------------------------------------------------------------------------------------
        # --------------------------BARS PLOT - HOUR OF A DAY -INTERNET--------------------------
        # -----------------(detailed internet connection count by hour of a day)-----------------
        # ---------------------------------------------------------------------------------------
    if "internet" in origin_type_list :
        working_days = [0, 1, 2, 3, 4]
        rest_days = [5, 6]
        temp_hours_plot_types_list = ['internet']

        # preparing a tab with WORKING days and active connections:
        df_internet_wd = df[df['type'].isin(temp_hours_plot_types_list) & df['day_of_week'].isin(working_days)].copy()
        # preparing a tab with REST days and active connections:
        df_internet_rd = df[df['type'].isin(temp_hours_plot_types_list) & df['day_of_week'].isin(rest_days)].copy()

        # get Series of WORKING days stat by hour:
        working_days_count = df_internet_wd['date'].nunique()
        con_by_hour_in_wd = df_internet_wd.groupby(['hour_of_day'])['weigth_column'].sum()
        con_by_hour_in_wd = con_by_hour_in_wd / working_days_count
        con_by_hour_in_wd.sort_index(inplace=True)
        con_by_hour_in_wd = con_by_hour_in_wd.reindex(range(24), fill_value=0)
        values_wd_by_hour = con_by_hour_in_wd.astype('float64')
        values_wd_by_hour = values_wd_by_hour.tolist()

        # get Series of REST days stat by hour:
        rest_days_count = df_internet_rd['date'].nunique()
        con_by_hour_in_rd = df_internet_rd.groupby(['hour_of_day'])['weigth_column'].sum()
        con_by_hour_in_rd = con_by_hour_in_rd / rest_days_count
        con_by_hour_in_rd.sort_index(inplace=True)
        con_by_hour_in_rd = con_by_hour_in_rd.reindex(range(24), fill_value=0)
        values_rd_by_hour = con_by_hour_in_rd.astype('float64')
        values_rd_by_hour = values_rd_by_hour.tolist()

        by_hour_x_ticks = ['00:00', '01:00', '02:00', '03:00', '04:00', '05:00', '06:00', '07:00', '08:00', '09:00',
                           '10:00', '11:00', '12:00', '13:00', '14:00', '15:00', '16:00', '17:00', '18:00', '19:00',
                           '20:00', '21:00', '22:00', '23:00']

        x = np.arange(len(by_hour_x_ticks))  # the label locations
        width = 0.3  # the width of the bars
        fig, ax = plt.subplots()
        rects1 = ax.bar(x - width / 2, values_wd_by_hour, width, label='У робочий день')
        rects2 = ax.bar(x + width / 2, values_rd_by_hour, width, label='У вихідний день')

        # Add some text for labels, title and custom x-axis tick labels, etc.
        ax.set_ylabel("ФІКСАЦІЙ НА ГОДИНУ")
        ax.set_title("СЕРЕДНЯ КІЛЬКІСТЬ ФІКСАЦІЙ МЕРЕЖІ ІНТЕРНЕТ У РОБОЧИЙ/ВИХІДНИЙ ДЕНЬ ПОГОДИННО\n ")
        ax.set_xticks(x)
        ax.set_xticklabels(by_hour_x_ticks)
        ax.legend()
        ax.grid(True)

        def autolabel(rects):
            """Attach a text label above each bar in *rects*, displaying its height."""
            for rect in rects:
                height = rect.get_height()
                ax.annotate('{:.1f}'.format(height),
                            xy=(rect.get_x() + rect.get_width() / 2, height),
                            xytext=(0, 3),  # 3 points vertical offset
                            textcoords="offset points",
                            ha='center', va='bottom', rotation='vertical')

        autolabel(rects1)
        autolabel(rects2)
        plt.xticks(rotation=45)
        fig.set_size_inches(15, 6, forward=True)
        fig.tight_layout()
        plt.savefig("by_hour_internet.png", bbox_inches='tight')
        pic_4 = document.add_paragraph(
            "Встановлені періоди доби в які абонент мав записи про використання мережі інтернет. "
            "На графіку надається середня "
            "кількість записів про фіксацію мережі інтернет у робочі та вихідні дні:"
            , style=body_text_style)
        document.add_paragraph('', style=body_text_style)
        pic_4.add_run().add_picture('by_hour_internet.png', width=Cm(17))
        document.add_paragraph('', style=body_text_style)
        plt.close(fig)

    # ---------------------------------------------------------------------------------------
    # -----------------------------------IMEI - GET STATISTIC--------------------------------
    # ---------------------------------------------------------------------------------------
    imei_series = df.groupby(['imei_a'])['weigth_column'].sum()
    imei_series.drop(labels=['', '0'], inplace=True, errors='ignore')
    imei_series.reset_index(drop=True)
    imei_list = imei_series.index.values.tolist()
    imei_values = imei_series.values.tolist()

    document.add_paragraph("Відомості щодо використаних пристрої зв’язку", style=main_header_style)
    document.add_paragraph("(періоди використання ІМЕІ):", style=main_header_style)
    document.add_paragraph("За період перевірки абонент використовував " + str(len(imei_list)) + " ІМЕІ:",
                           style=body_text_style)

    def imei_get_info(imei, value):
        df_for_imei = df.loc[df['imei_a'] == imei]
        imei_start = df_for_imei['date_time'].min()
        imei_start_date = mpl_dates.date2num(imei_start)
        imei_start_str = imei_start.strftime('%d.%m.%Y (з %H:%M год)')
        imei_end = df_for_imei['date_time'].max()
        imei_end_date = mpl_dates.date2num(imei_end)
        imei_end_str = imei_end.strftime('%d.%m.%Y (до %H:%M год)')
        value = str(value)
        imei = str(imei)
        ans = [imei, value, imei_start_str, imei_end_str, imei_start_date, imei_end_date]
        return ans

    imeis_all_info = []

    for i in range(len(imei_list)):
        imei_pack = imei_get_info(imei_list[i], imei_values[i])
        document.add_paragraph( str(imei_pack[0]) + " всього записів про з'єднання - " + str(imei_pack[1])
                                + ", у період часу: " + str(imei_pack[2]) + " - " + str(imei_pack[3]),
                                style=list_style)
        imeis_all_info.append(imei_pack)

    get_1day_val = mpl_dates.date2num(date_start_tab + pd.Timedelta('1d')) - mpl_dates.date2num(date_start_tab)

    def imei_add_dates_array(imeis_info):
        df_for_imei = df.loc[df['imei_a'] == imeis_info[0]]
        dates_list = df_for_imei['date_time'].copy()
        dates_list = dates_list.reset_index(drop=True)
        second_date_imei = get_1day_val
        result_imeis_periods = []
        for d in range(dates_list.size):
            first_date_imei = mpl_dates.date2num(dates_list.at[d])
            result_imeis_periods.append([first_date_imei, second_date_imei])
            imeis_info.append(result_imeis_periods)
        return imeis_info

    imeis_plot_info = []
    for f in range(len(imeis_all_info)):
        imeis_plot_info.append(imei_add_dates_array(imeis_all_info[f]))
    # so now we have array 'imeis_all_info' =
    # [[ [0]imei, [1]value, [2]start_str, [3]end_str, [4]start_date, [5]end_date, [6] array of usage [date, dur] ]]

    fig, ax = plt.subplots()
    fig.set_size_inches(15, 1*(len(imeis_all_info)), forward=True)
    ax.xaxis.set_major_formatter(mpl_dates.DateFormatter("%d.%m.%Y"))
    plt.locator_params(axis='x', nbins=20)
    plt.xticks(rotation=45)
    for p in range(len(imeis_all_info)):
        ax.broken_barh(imeis_all_info[p][6], (((p+1)*10)+1, 8), facecolors='tab:red')

    start_imei_plot = mpl_dates.date2num(date_start_tab)
    end_imei_plot = mpl_dates.date2num(date_end_tab)

    ax.set_ylim(10, (len(imeis_all_info)+1)*10)
    ax.set_xlim(start_imei_plot, end_imei_plot)
    ax.set_title("\nПЕРІОДИ ФІКСАЦІЇ ІМЕІ\n ")

    y_ticks_labels_imei = []
    y_ticks_pos_imei = []
    for numbers in range(len(imeis_all_info)):
        y_ticks_labels_imei.append(imeis_all_info[numbers][0])
        y_ticks_pos_imei.append(15+(numbers*10))
    ax.set_yticks(y_ticks_pos_imei)

    ax.set_yticklabels(y_ticks_labels_imei)
    ax.grid(True)
    plt.savefig("imei_broken_barh.png", bbox_inches='tight')

    pic_5 = document.add_paragraph('', style=main_header_style)
    pic_5.add_run().add_picture('imei_broken_barh.png', width=Cm(17))
    plt.close(fig)


    # ---------------------------------------------------------------------------------------
    # -----------------------------------BS - statistic--------------------------------------
    # ---------------------------------------------------------------------------------------
    document.add_paragraph("Відомості щодо використання абонентом", style=main_header_style)
    document.add_paragraph("базових станцій зв’язку (БС):", style=main_header_style)

    bs_series = df.groupby(['adr_a'])['weigth_column'].sum()
    bs_series.drop(labels=['', '0'], inplace=True, errors='ignore')
    sum_all_bs_values = bs_series.sum()
    bs_name_count = bs_series.size
    select_level_value = (float(sum_all_bs_values) / 100)*2

    document.add_paragraph("За період перевірки зафіксовано використання абонентом БС у кількості -"
                           + str(bs_name_count)
                           + " (найменувань адрес). У вибірку додані адреси БС, якими абонент користувався "
                             "найчастіше > 1%):",
                           style=body_text_style)

    bs_series = bs_series[bs_series > select_level_value]
    bs_series.sort_values(inplace=True, ascending=False)
    bs_top_list = bs_series.index.values.tolist()
    bs_top_values = bs_series.values.tolist()

    def get_frequent_azimuth(bs_name):
        df_azimuth_get = df.loc[df['adr_a'] == bs_name].copy()
        azimuth_rating = df_azimuth_get.groupby(['az_a'])['weigth_column'].sum()
        azimuth_rating.drop(labels=['', '0'], inplace=True, errors='ignore')
        if azimuth_rating.size > 0:
            list_az_names = azimuth_rating.index.values.tolist()
            azimuth = list_az_names[0]
        else:
            azimuth = 'у первинній таблиці відсутні'
        return azimuth

    azimuth_fr_list = []
    for bs in bs_top_list:
        azimuth_fr_list.append(get_frequent_azimuth(bs))

    for num in range(len(bs_top_list)):
        document.add_paragraph(str(bs_top_values[num]) + " - разів зафіксовано: найменування адреси розташування "
                                                         "базової станції: " + str(bs_top_list[num]) +
                               " (азимут найчастіше - " + str(azimuth_fr_list[num]) + ").",
                               style=list_style)

    def build_plots_for_bs(index_bs):
        document.add_paragraph(str(index_bs+1) + ". Відомості про використання БС:", style=main_header_style)
        document.add_paragraph(str(bs_top_list[index_bs]) + ", азимут = " + str(azimuth_fr_list[index_bs]) +
                               " (" + str(bs_top_values[index_bs]) + " записів)",
                               style=main_header_style)

        df_bs_usage = df.loc[df['adr_a'].isin([bs_top_list[index_bs]])].copy()
        series_bs_usage = df_bs_usage.groupby(['date'])['weigth_column'].sum()
        values_bs_usage = series_bs_usage.astype('int64')
        values_bs_usage = values_bs_usage.tolist()
        date_bs_usage = series_bs_usage.index.values.tolist()
        for item2 in range(len(date_bs_usage)):
            date_bs_usage[item2] = datetime.strptime(date_bs_usage[item2], '%d.%m.%Y')
        date_bs_usage = mpl_dates.date2num(date_bs_usage)

        fig6, ax6 = plt.subplots()
        fig6.set_size_inches(18.5, 4, forward=True)
        ax6.xaxis.set_major_formatter(mpl_dates.DateFormatter("%d.%m.%Y"))
        plt.locator_params(axis='y', nbins=10)
        plt.locator_params(axis='x', nbins=25)
        ax6.set_xlim(plot_start, plot_end)
        lim_y_bs = (sum(values_bs_usage)/len(date_bs_usage))*3
        ax6.set_ylim(0, lim_y_bs)
        plt.xticks(rotation=45)
        ax6.grid(True)
        ax6.set(xlabel=str(bs_top_list[index_bs]),
                ylabel="ВИКОРИСТАННЯ БАЗОВОЇ СТАНЦІЇ\n (кількість записів на день)",
                title="ІНТЕНСИВНІСТЬ ВИКОРИСТАННЯ БС")
        ax6.bar(date_bs_usage, values_bs_usage, color='blue')
        plt.savefig(str(subscriber) + "_bs_timeline_" + str(index_bs) + ".png", bbox_inches='tight')
        document.add_picture(str(subscriber) + "_bs_timeline_" + str(index_bs) + ".png", width=Cm(17))
        plt.close(fig6)

        working_days_bs = [0, 1, 2, 3, 4]
        rest_days_bs = [5, 6]
        bs_to_search = [str(bs_top_list[index_bs]), 'ytn nfrjq']

        # preparing a tab with WORKING days and active connections:
        df_bs_wd = df[df['adr_a'] .isin(bs_to_search) & df['day_of_week'].isin(working_days_bs)].copy()
        # preparing a tab with REST days and active connections:
        df_bs_rd = df[df['adr_a'].isin(bs_to_search) & df['day_of_week'].isin(rest_days_bs)].copy()

        # get Series of WORKING days stat by hour:
        working_days_count_bs = df_bs_wd['date'].nunique()
        bs_by_hour_in_wd = df_bs_wd.groupby(['hour_of_day'])['weigth_column'].sum()
        bs_by_hour_in_wd = bs_by_hour_in_wd / working_days_count_bs
        bs_by_hour_in_wd.sort_index(inplace=True)
        bs_by_hour_in_wd = bs_by_hour_in_wd.reindex(range(24), fill_value=0)
        values_wd_by_hour_bs = bs_by_hour_in_wd.astype('float64')
        values_wd_by_hour_bs = values_wd_by_hour_bs.tolist()

        # get Series of REST days stat by hour:
        rest_days_count_bs = df_bs_rd['date'].nunique()
        bs_by_hour_in_rd = df_bs_rd.groupby(['hour_of_day'])['weigth_column'].sum()
        bs_by_hour_in_rd = bs_by_hour_in_rd / rest_days_count_bs
        bs_by_hour_in_rd.sort_index(inplace=True)
        bs_by_hour_in_rd = bs_by_hour_in_rd.reindex(range(24), fill_value=0)
        values_rd_by_hour_bs = bs_by_hour_in_rd.astype('float64')
        values_rd_by_hour_bs = values_rd_by_hour_bs.tolist()

        by_hour_x_ticks_bs = ['00:00', '01:00', '02:00', '03:00', '04:00', '05:00', '06:00', '07:00', '08:00', '09:00',
                           '10:00', '11:00', '12:00', '13:00', '14:00', '15:00', '16:00', '17:00', '18:00', '19:00',
                           '20:00', '21:00', '22:00', '23:00']

        x1 = np.arange(len(by_hour_x_ticks_bs))  # the label locations
        width1 = 0.3  # the width of the bars
        fig8, ax8 = plt.subplots()
        rects11 = ax8.bar(x1 - width1 / 2, values_wd_by_hour_bs, width1, label='У робочий день')
        rects22 = ax8.bar(x1 + width1 / 2, values_rd_by_hour_bs, width1, label='У вихідний день')

        # Add some text for labels, title and custom x-axis tick labels, etc.
        ax8.set_ylabel("ФІКСАЦІЙ НА ГОДИНУ")
        ax8.set_xlabel(str(bs_top_list[index_bs]))
        ax8.set_title("СЕРЕДНЯ КІЛЬКІСТЬ ФІКСАЦІЙ БС У РОБОЧИЙ/ВИХІДНИЙ ДЕНЬ ПОГОДИННО\n ")
        ax8.set_xticks(x1)
        ax8.set_xticklabels(by_hour_x_ticks_bs)
        ax8.legend()
        ax8.grid(True)

        def autolabel1(rects):
            """Attach a text label above each bar in *rects*, displaying its height."""
            for rect in rects:
                height = rect.get_height()
                ax.annotate('{:.1f}'.format(height),
                            xy=(rect.get_x() + rect.get_width() / 2, height),
                            xytext=(0, 3),  # 3 points vertical offset
                            textcoords="offset points",
                            ha='center', va='bottom', rotation='vertical')

        autolabel1(rects11)
        autolabel1(rects22)
        plt.xticks(rotation=45)
        fig8.set_size_inches(15, 6, forward=True)
        fig8.tight_layout()
        plt.savefig(str(subscriber) + "_bs_daybar_" + str(index_bs) + ".png", bbox_inches='tight')

        document.add_picture(str(subscriber) + "_bs_daybar_" + str(index_bs) + ".png", width=Cm(17))
        document.add_paragraph('', style=body_text_style)
        plt.close(fig)

    for bs_ploting in range(len(bs_top_list)):
        build_plots_for_bs(bs_ploting)

    # ---------------------------------------------------------------------------------------
    # ---------------------------------SUBSCRIBER ANALYSIS-----------------------------------
    # ---------------------------------------------------------------------------------------
    document.add_paragraph("Відомості щодо найближчого оточення фігуранта", style=main_header_style)
    document.add_paragraph("(абоненти з найбільшою кількістю зв'язків):", style=main_header_style)

    df_fr_all = df[df['type'].isin(["вх", "вих", "вх СМС", "вих СМС", "переад"])].copy()
    fr_series = df_fr_all.groupby(['sim_b'])['weigth_column'].sum()
    fr_series.drop(labels=['', '0', 'KYIVSTAR', 'LIFECELL', 'lifecell', '6969', '5050', 'lifecell.ua', '5433', '2080', 'MO3', '693', 'VODAFONE'], inplace=True, errors='ignore')
    fr_series.sort_values(inplace=True, ascending=False)
    sum_all_fr_values = fr_series.sum()
    fr_name_count = fr_series.size
    select_level_value_fr = (float(sum_all_fr_values) / 100)*2

    document.add_paragraph("Перевіркою кількості контактів з іншими абонентами виявлено співрозмовників, "
                           "що є найближчими до фігуранта (абонентів, що мають щонайменше 2% питомої ваги у "
                           "спілкуванні фігуранта):", style=body_text_style)

    fr_series = fr_series[fr_series > select_level_value_fr]
    fr_series.sort_values(inplace=True, ascending=False)
    fr_top_list = fr_series.index.values.tolist()
    fr_top_values = fr_series.values.tolist()

    def get_fr_info(fr_pos):
        fr_now = str(fr_top_list[fr_pos])
        df_friend = df_fr_all.loc[df_fr_all['sim_b'] == fr_now].copy()
        # append main info [0] - subscriber, [1] - count connections total
        one_fr_info = []
        one_fr_info.append(fr_top_list[fr_pos])
        one_fr_info.append(fr_top_values[fr_pos])

        # append types stat [2] - used types, [3] - count types:
        types_fr_rating = df_friend.groupby(['types_stat'])['weigth_column'].sum()
        types_fr_rating.drop(labels=['', '0'], inplace=True, errors='ignore')
        list_fr_types_names = types_fr_rating.index.values.tolist()
        list_fr_types_vals = types_fr_rating.values.tolist()
        one_fr_info.append(list_fr_types_names)
        one_fr_info.append(list_fr_types_vals)

        # get dur in and out values [4] - dur in, [5] - dur out :
        dur_fr_all_types = df_friend.groupby(['type'])['dur'].sum()
        if 'вх' in dur_fr_all_types:
            dur_in_fr = dur_fr_all_types.at['вх']
        else:
            dur_in_fr = pd.Timedelta('0s')
        if 'вих' in dur_fr_all_types:
            dur_out_fr = dur_fr_all_types.at['вих']
        else:
            dur_out_fr = pd.Timedelta('0s')
        one_fr_info.append(dur_in_fr)
        one_fr_info.append(dur_out_fr)

        # get connections count by date [6] - dates, [7] - connections for each day:
        con_fr_by_date = df_friend.groupby(['date'])['weigth_column'].sum()
        values_fr_by_date = con_fr_by_date.astype('int64')
        values_fr_by_date = values_fr_by_date.tolist()
        date_fr_by_date = con_fr_by_date.index.values.tolist()
        for fr_date in range(len(date_fr_by_date)):
            date_fr_by_date[fr_date] = datetime.strptime(date_fr_by_date[fr_date], '%d.%m.%Y')
        date_fr_by_date = mpl_dates.date2num(date_fr_by_date)
        one_fr_info.append(date_fr_by_date)
        one_fr_info.append(values_fr_by_date)

        # get count by hour of a day [8] - working days, [9] - rest days
        working_days_fr = [0, 1, 2, 3, 4]
        rest_days_fr = [5, 6]
        fr_to_search = [str(fr_top_list[fr_pos]), 'only need second element']
        # preparing a tab with WORKING days and active connections:
        df_fr_wd = df[df['sim_b'] .isin(fr_to_search) & df['day_of_week'].isin(working_days_fr)].copy()
        # preparing a tab with REST days and active connections:
        df_fr_rd = df[df['sim_b'].isin(fr_to_search) & df['day_of_week'].isin(rest_days_fr)].copy()
        # get Series of WORKING days stat by hour:
        working_days_count_fr = df_fr_wd['date'].nunique()
        fr_by_hour_in_wd = df_fr_wd.groupby(['hour_of_day'])['weigth_column'].sum()
        fr_by_hour_in_wd = fr_by_hour_in_wd / working_days_count_fr
        fr_by_hour_in_wd.sort_index(inplace=True)
        fr_by_hour_in_wd = fr_by_hour_in_wd.reindex(range(24), fill_value=0)
        values_wd_by_hour_fr = fr_by_hour_in_wd.astype('float64')
        values_wd_by_hour_fr = values_wd_by_hour_fr.tolist()
        # get Series of REST days stat by hour:
        rest_days_count_fr = df_fr_rd['date'].nunique()
        fr_by_hour_in_rd = df_fr_rd.groupby(['hour_of_day'])['weigth_column'].sum()
        fr_by_hour_in_rd = fr_by_hour_in_rd / rest_days_count_fr
        fr_by_hour_in_rd.sort_index(inplace=True)
        fr_by_hour_in_rd = fr_by_hour_in_rd.reindex(range(24), fill_value=0)
        values_rd_by_hour_fr = fr_by_hour_in_rd.astype('float64')
        values_rd_by_hour_bs = values_rd_by_hour_fr.tolist()
        one_fr_info.append(values_wd_by_hour_fr)
        one_fr_info.append(values_rd_by_hour_bs)
        one_fr_info.append(fr_pos)
        return one_fr_info

    fr_main_value = []
    for pos_friend in range(len(fr_top_list)):
        fr_main_value.append(get_fr_info(pos_friend))

    def subscriber_analysis(friend):
        fr_dur_lits = [friend[5].components.hours * 3600 + friend[5].components.minutes * 60
                       + friend[5].components.seconds,
                       friend[4].components.hours * 3600 + friend[4].components.minutes * 60
                       + friend[4].components.seconds]
        all_dur = friend[5] + friend[4]
        all_dur_str = f'{all_dur.components.hours:02d}:{all_dur.components.minutes:02d}:{all_dur.components.seconds:02d} сек.'
        document.add_paragraph(str(friend[10]+1) + ". Співрозмовник " + str(friend[0]), style=main_header_style)
        document.add_paragraph("Фігурант має з'єднання з абонентом " + str(friend[0]) + " в кількості - "
                               + str(friend[1]) + " записів, загальна тривалість розмов складає "
                               + all_dur_str + ". Абоненти мають наступні звички у спілкуванні:", style=body_text_style)
        fr_out_dur_str = f'{friend[5].components.hours:02d}:{friend[5].components.minutes:02d}:{friend[5].components.seconds:02d} сек.'
        fr_in_dur_str = f'{friend[4].components.hours:02d}:{friend[4].components.minutes:02d}:{friend[4].components.seconds:02d} сек.'
        # drawing types statistic plots:
        types_labels = friend[2]
        plt.rcdefaults()
        plt.rcParams['font.size'] = 18.0
        fig1, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 4))
        y_pos = np.arange(len(types_labels))
        ax1.barh(types_labels, friend[3], color='blue')
        ax1.set_yticks(y_pos)
        ax1.set(ylabel="ТИПИ З'ЄДНАНЬ")
        ax1.set(xlabel="КІЛЬКІСТЬ \n(всього за період перевірки - " + str(friend[1]) + ")")
        ax1.invert_yaxis()
        pie_type_list_fr = ['Вихідні розмови:\n' + fr_out_dur_str, 'Вхідні розмови:\n' + fr_in_dur_str]
        ax2.pie(fr_dur_lits, labels=pie_type_list_fr, startangle=0, autopct='%0.1F%%', shadow=True,
                           wedgeprops={"edgecolor": "k", 'linewidth': 1}, labeldistance=1.2, colors=['cyan', 'magenta'],
                           explode=[0.1, 0])
        ax2.text(-1.25, -1.7, 'РОЗМОВ ВСЬОГО: ' + all_dur_str)
        plt.savefig(str(subscriber) + "_" + str(friend[0]) + "1st.png", bbox_inches='tight')
        document.add_paragraph('', style=body_text_style)
        document.add_picture(str(subscriber) + "_" + str(friend[0]) + "1st.png", width=Cm(17))
        document.add_paragraph('', style=body_text_style)
        plt.close(fig1)

        fig6, ax6 = plt.subplots()
        fig6.set_size_inches(18.5, 4, forward=True)
        ax6.xaxis.set_major_formatter(mpl_dates.DateFormatter("%d.%m.%Y"))
        plt.locator_params(axis='y', nbins=10)
        plt.locator_params(axis='x', nbins=25)
        ax6.set_xlim(plot_start, plot_end)
        lim_y_fr = (sum(friend[7])/len(friend[6]))*3
        ax6.set_ylim(0, lim_y_fr)
        plt.xticks(rotation=45)
        ax6.grid(True)
        ax6.set(ylabel="КОНТАКТИ МІЖ АБОНЕНТАМИ\n (кількість записів на день)",
                title="ПЕРІОДИ СПІЛКУВАННЯ З" + str(friend[0]))
        ax6.bar(friend[6], friend[7], color='blue')
        plt.savefig(str(subscriber) + "_" + str(friend[0]) + "_2nd.png", bbox_inches='tight')
        document.add_picture(str(subscriber) + "_" + str(friend[0]) + "_2nd.png", width=Cm(17))
        plt.close(fig6)

        by_hour_x_ticks_fr = ['00:00', '01:00', '02:00', '03:00', '04:00', '05:00', '06:00', '07:00', '08:00', '09:00',
                           '10:00', '11:00', '12:00', '13:00', '14:00', '15:00', '16:00', '17:00', '18:00', '19:00',
                           '20:00', '21:00', '22:00', '23:00']

        x1 = np.arange(len(by_hour_x_ticks_fr))  # the label locations
        width1 = 0.3  # the width of the bars
        fig8, ax8 = plt.subplots()
        rects11 = ax8.bar(x1 - width1 / 2, friend[8], width1, label='У робочий день')
        rects22 = ax8.bar(x1 + width1 / 2, friend[9], width1, label='У вихідний день')

        # Add some text for labels, title and custom x-axis tick labels, etc.
        ax8.set_ylabel("ФІКСАЦІЙ НА ГОДИНУ")
        # ax8.set_xlabel(str(bs_top_list[index_bs]))
        ax8.set_title("СЕРЕДНЯ КІЛЬКІСТЬ КОНТАКТІВ У РОБОЧИЙ/ВИХІДНИЙ ДЕНЬ ПОГОДИННО\n ")
        ax8.set_xticks(x1)
        ax8.set_xticklabels(by_hour_x_ticks_fr)
        ax8.legend()
        ax8.grid(True)

        def autolabel1(rects):
            """Attach a text label above each bar in *rects*, displaying its height."""
            for rect in rects:
                height = rect.get_height()
                ax.annotate('{:.1f}'.format(height),
                            xy=(rect.get_x() + rect.get_width() / 2, height),
                            xytext=(0, 3),  # 3 points vertical offset
                            textcoords="offset points",
                            ha='center', va='bottom', rotation='vertical')

        autolabel1(rects11)
        autolabel1(rects22)
        plt.xticks(rotation=45)
        fig8.set_size_inches(15, 6, forward=True)
        fig8.tight_layout()
        plt.savefig(str(subscriber) + "_" + str(friend[0]) + "_3d.png", bbox_inches='tight')

        document.add_picture(str(subscriber) + "_" + str(friend[0]) + "_3d.png", width=Cm(17))
        document.add_paragraph('', style=body_text_style)
        plt.close(fig)

    for fr2anal in range(len(fr_main_value)):
        subscriber_analysis(fr_main_value[fr2anal])


    document.save('analysis_' + str(subscriber) + '.docx')





